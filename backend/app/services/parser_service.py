import io
from pypdf import PdfReader
from docx import Document

class ParserService:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    @classmethod
    def parse_file(cls, filename: str, file_bytes: bytes) -> str:
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return cls.extract_text_from_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return cls.extract_text_from_docx(file_bytes)
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
